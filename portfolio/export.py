import frappe
from frappe.utils.pdf import get_pdf
from frappe import _
import io
import requests
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt
from datetime import datetime


@frappe.whitelist()
def export_portfolio(portfolio_names, format):
	if not portfolio_names:
		frappe.throw(_("No portfolio names provided"))

	file_data_list = []
	content = generate_html_content(portfolio_names)

	if format == "pdf":
		file_data = get_pdf(content)
		file_extension = "pdf"
	elif format == "docx":
		file_data = generate_docx(content)
		file_extension = "docx"
	elif format == "world_bank":
		file_data = worldbank_format(portfolio_names)
		file_extension = "docx"
	else:
		frappe.throw(_("Unsupported file format"))

	file_data_list.append(file_data)

	# Combine file_data_list into a single file if needed, for simplicity, return the first file.
	combined_file_data = file_data_list[0]

	# Generate a default file name with timestamp
	timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
	default_file_name = f"portfolio_export_{timestamp}.{file_extension}"

	# Create a new File document to save the generated file
	file_doc = frappe.get_doc({
		"doctype": "File",
		"file_name": default_file_name,
		"is_private": 1,
		"content": combined_file_data
	})
	file_doc.insert()

	return {
		"status": "success",
		"message": f"Portfolios exported successfully.",
		"file_url": file_doc.file_url
	}


def generate_html_content(portfolios):
	project_details = ""
	portfolio_names = frappe.parse_json(portfolios)
	for docname in portfolio_names:
		portfolio = frappe.get_doc("Portfolio", docname)
		technologies_list = ""
		for tech in portfolio.technologies:
			technologies_list += f"<li>{tech.technology_name}</li>"

		images_list = ""
		for image in portfolio.images:
			if image:
				images_list += f'<img src="{image.website_image}" alt="Screenshot" style="width:300px;height:auto;"><br>'

		project_details += f"""
		<h3>{portfolio.title}</h3>
		{images_list}
		<p>Client: {portfolio.client}</p>
		<p>Period: {portfolio.start_date} - {portfolio.end_date}</p>
		<p>Technologies:</p>
		<ul>
		{technologies_list}
		</ul>
		<p>Details: {portfolio.body}</p>
		<p>URL: {portfolio.website}</p>
		<p>Location: {portfolio.location}</p>
		<hr>
		"""

	content = f"""
    <html>
    <head>
        <title>Portfolio: PAST AND CURRENT PROJECTS</title>
    </head>
    <body>
        <h1>PAST AND CURRENT PROJECTS</h1>
        <p>Here is a selection of projects we have worked on or are working on.</p>
        {project_details}
    </body>
    </html>
    """
	return content


def generate_docx(html_content):
	doc = Document()
	soup = BeautifulSoup(html_content, 'html.parser')

	for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'li', 'img']):
		if element.name.startswith('h'):
			doc.add_heading(element.get_text(), level=int(element.name[1]))
		elif element.name == 'p':
			doc.add_paragraph(element.get_text())
		elif element.name == 'ul':
			for li in element.find_all('li'):
				doc.add_paragraph(f'• {li.get_text()}', style='ListBullet')
		elif element.name == 'img':
			response = requests.get(element['src'])
			img = io.BytesIO(response.content)
			doc.add_picture(img, width=Inches(4.0))

	output = io.BytesIO()
	doc.save(output)
	return output.getvalue()


def worldbank_format(portfolios):
	"""Create a world bank format."""
	portfolio_names = frappe.parse_json(portfolios)
	doc = Document()

	# Add Title
	title = doc.add_heading(level=1)
	title_run = title.add_run("Assignment Details")
	title_run.bold = True

	# Loop through each portfolio and create a table
	for portfolio_name in portfolio_names:
		details = frappe.get_doc("Portfolio", portfolio_name)

		# Add a heading for each portfolio
		doc.add_heading(portfolio_name, level=2)

		# Create a table for the details
		table = doc.add_table(rows=14, cols=2)
		table.style = 'Table Grid'
		table.autofit = False

		# Set the width of the table columns
		for row in table.rows:
			row.cells[0].width = Pt(200)
			row.cells[1].width = Pt(300)

		# Add the details to the table
		details_dict = {
			"Assignment name:": details.title,
			"Approx. value of the contract (in current US$):": details.approximate_contract_value,
			"Country:": details.location,
			"Duration of assignment (months):": details.duration_of_assignment,
			"Name of Client(s):": details.project.client,
			"Contact Person, Title/Designation, Tel. No./Address:": details.contact,
			"Start Date (month/year):": details.start_date,
			"End Date (month/year):": details.end_date,
			"Total No. of staff-months of the assignment:": details.total_staff_months,
			"No. of professional staff-months provided by your consulting firm/organization or "
			"your sub consultants:": details.total_staff_months,
			"Name of associated Consultants, if any:": "",
			"Name of senior professional staff of your consulting firm/organization involved and "
			"designation and/or functions performed (e.g. Project Director/ Coordinator, "
			"Team Leader):": "",
			"Description of Project:": details.body,
			"Description of actual services provided by your staff within the "
			"assignment:": details.serservices_listed,
		}

		# Populate the table
		for i, (key, value) in enumerate(details_dict.items()):
			table.cell(i, 0).text = key
			table.cell(i, 1).text = value

	# Save the document
	doc_path = "/mnt/data/worldbank_format.docx"
	doc.save(doc_path)
	print("Document saved as worldbank_format.docx")
	return doc_path

# Generate the document
